import sys

# Pre-flight check to ensure the library is installed
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("CRITICAL ERROR: The python-docx library is missing.")
    print("Please run this command in your terminal first:")
    print("pip install python-docx")
    sys.exit(1)

doc = Document()

# --- GLOBAL STYLE CONFIGURATION ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# --- HELPER FUNCTIONS FOR ACADEMIC FORMATTING ---
def add_h(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True

def add_p(text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    if bold:
        r.bold = True

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.add_run(text)

def add_image_placeholder(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(f"\n[ INSERT IMAGE HERE: {text} ]\n")
    r.bold = True
    r.font.color.rgb = RGBColor(255, 0, 0) # Red color so you don't miss it

print("Generating Document...")

# --- TITLE PAGE ---
doc.add_paragraph("\n\n\n\n")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("UPSCALE PRO\n")
r.bold = True
r.font.size = Pt(24)
r2 = p.add_run("MAJOR PROJECT / DISSERTATION REPORT\n\n\n")
r2.font.size = Pt(16)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run("Submitted by: Amir Shamim (2023-301-021)\n").bold = True
p2.add_run("In partial fulfilment for the award of the degree of:\nBACHELOR OF COMPUTER APPLICATION\n\n")
p2.add_run("Under the supervision of: Dr. Sapna Jain Ma’am\n\n").bold = True
p2.add_run("Institution: Jamia Hamdard University, Department of Computer Science & Technology\n")
doc.add_page_break()

# --- CONTENT GENERATION ---

add_h("DECLARATION", 1)
add_p("I, Amir Shamim, a student of BCA (Enrollment No: 2023-301-021), hereby declare that the dissertation entitled “Upscale Pro” which is being submitted by me to the Department of Computer Science, Jamia Hamdard, New Delhi in partial fulfilment of the requirement for the award of the degree of BCA is my original work and has not been submitted anywhere else for the award of any Degree, Diploma, Associateship, Fellowship or other similar title or recognition.")
p = doc.add_paragraph()
p.add_run("\nApril, 2026\n").bold = True
p.add_run("Amir Shamim").bold = True
doc.add_page_break()

add_h("ACKNOWLEDGEMENT", 1)
add_p("It is a pleasure to acknowledge many people who knowingly and unknowingly helped me to complete my project. First, I thank God for all the blessings that carried me through all these years.")
add_p("I extend my utmost gratitude to Dr Sapna Jain Ma’am, project supervisor, who has always supported, guided, appreciated, and encouraged me to explore more opportunities. She enlightened me at various stages during the development of this project and provided many insights and useful examples, which proved immensely helpful in its successful completion.")
add_p("I extend my sincere gratitude to my teachers and guides who made unforgettable contributions. I thank all the non-teaching staff of our institution who were always ready to help in whatever way they could.")
doc.add_page_break()

add_h("1. OBJECTIVE", 1)
add_p("The main goal of this project is to build and launch a scalable web platform that uses Generative Adversarial Networks (GANs)—specifically, highly tuned versions of Real-ESRGAN—to upscale images easily. I want to make advanced AI upscaling, which usually requires high-end local hardware, accessible to everyday users through a simple, browser-based web interface.")
add_p("The specific objectives are:")
add_bullet("To design and build a full-stack SaaS (Software as a Service) ecosystem. This utilizes React.js for a responsive client frontend, Node.js for a robust backend API gateway, and a decoupled serverless architecture to handle intensive GPU processing without bottlenecking the main server.")
add_bullet("To cut down processing time to a reliable sub-10-second benchmark. This is achieved by offloading the heavy PyTorch inference tasks to Modal.com's serverless GPUs. Traditional cloud GPUs require expensive continuous uptime, but this architecture utilizes ephemeral containers optimized to minimize cold-start delays, making the application both fast and financially viable.")
add_bullet("To set up a secure, rate-limited system that can handle multiple concurrent users. The system must process simultaneous image uploads and queue them intelligently without crashing the primary Node.js server.")
add_bullet("To offer tiered processing models (such as Real-ESRGAN Pro for standard images and Real-ESRGAN Anime for digital illustrations) based on user subscription levels, managing all authentication and billing data securely using a PostgreSQL database.")
doc.add_page_break()

add_h("2. INTRODUCTION", 1)
add_p("Today, digital images dictate how we interact with the web—from online shopping catalogs and social media platforms to digital archiving and professional broadcasting. But getting high-quality, high-resolution pictures isn't always possible. Hardware constraints on older mobile devices, aggressive compression algorithms used by messaging apps, and limited bandwidth often leave users with low-quality, heavily pixelated files.")
add_p("This project introduces a new, serverless GPU-powered Image Upscaler built as a commercial SaaS. It is specifically designed to take low-resolution images and automatically rebuild them into sharp, high-resolution versions. By shifting the computational burden from the user's local machine to cloud-based neural networks, the platform democratizes access to professional-grade image enhancement without requiring the user to own expensive graphics cards.")
add_h("Evolution of Scaling Algorithms", 2)
add_p("Before deep learning and neural networks became the industry standard, developers mainly relied on basic spatial domain algorithms to resize images. Traditional mathematical methods like nearest-neighbor, bilinear, and bicubic interpolation simply calculate and guess the missing pixels based on the color values of the pixels immediately surrounding them. While these methods are computationally cheap and extremely fast, they fail to understand the actual content of the image. Consequently, they often make images look blocky or severely blurred, permanently ruining fine high-frequency details like skin textures, fabric threads, and sharp geometric edges.")
add_p("To fix this fundamental flaw, researchers introduced machine learning to the pipeline. Early Convolutional Neural Networks (CNNs) like SRCNN learned how to map low-resolution images directly to high-resolution outputs by training on thousands of image pairs. However, the real breakthrough in perceptual quality came with Generative Adversarial Networks (GANs). Models like ESRGAN set up a competitive environment between two distinct networks: a Generator that tries to create realistic details, and a Discriminator that tries to spot the fake, AI-generated pixels. The Real-ESRGAN model used in this project takes this architecture a step further. It utilizes a complex \"degradation model\" that simulates real-world image problems—like motion blur, heavy JPEG compression, and camera sensor noise—so it can effectively clean up and sharpen dirty, everyday photos rather than just pristine laboratory samples.")
doc.add_page_break()

add_h("3. PROBLEM STATEMENT", 1)
add_p("Low-resolution images inherently lack fine spatial details. When these images are displayed on modern 4K monitors or high-density smartphone screens (like Retina displays), the lack of pixel density becomes obvious, resulting in a pixelated and heavily blurred viewing experience. Older mathematical resizing methods cannot synthesize this missing data; they just blur the existing pixels together, which often makes the picture look significantly worse.")
add_p("The core problem this project addresses is twofold:")
add_bullet("Traditional algorithms cannot create perceptually realistic high-resolution images.")
add_bullet("Running the advanced AI models that can hallucinate these missing details requires massive computational power, specifically from Graphics Processing Units (GPUs). Historically, if a developer tried to run these heavy PyTorch models on standard monolithic web servers (like a basic VPS), the entire system would freeze up during processing. This architectural flaw leads to extremely slow loading times, massive server hosting bills, and inevitable system timeouts the moment multiple people attempt to use the application simultaneously.")
doc.add_page_break()

add_h("4. SYSTEM REQUIREMENT SPECIFICATION", 1)
add_h("4.1 Hardware Requirements", 2)
add_bullet("Client Side: A standard consumer-grade PC or mobile device equipped with a modern web browser (Chrome, Firefox, Safari). No dedicated GPU is required on the user's end.")
add_bullet("Server Side Layer: Lightweight, cost-effective cloud servers (such as DigitalOcean Droplets) to host the Node.js API and manage incoming web traffic.")
add_bullet("Inference Hardware: On-demand, serverless NVIDIA T4 or A10G GPUs provided by Modal.com. These are used exclusively for rapid AI processing and shut down immediately when not in use.")
add_h("4.2 Software Technology Stack", 2)
add_p("This project combines the speed and agility of a modern JavaScript web frontend with a strictly isolated Python/PyTorch backend environment for the heavy lifting.")
add_bullet("Frontend Stack: React.js (bundled with Vite for fast Hot Module Replacement and optimized build times). The UI features a Virtual DOM implementation to offer rapid, interactive side-by-side visual image comparisons without page reloads.")
add_bullet("Backend Server: Node.js encapsulating an Express.js API gateway. This server acts as the traffic controller—it handles user authentication, securely sanitizes incoming file payloads, validates JWTs (JSON Web Tokens), and pushes jobs to the processing queue.")
add_bullet("Database Architecture (Hybrid): The system utilizes a dual-database strategy. It leverages SQLite for rapid, localized development and seamlessly switches to Neon Serverless PostgreSQL for production. This architecture ensures native connection pooling, making it highly reliable for storing user account details, dynamic subscription tiers, and robust usage analytics.")
add_bullet("Storage (CDN): Cloudinary cloud storage arrays. This acts as a Content Delivery Network (CDN) to safely store and quickly serve the finalized high-resolution images back to the user.")
add_bullet("AI Processing: Modal serverless Python containers running the PyTorch machine learning library and specific Real-ESRGAN model weights.")
add_h("4.3 Functional Requirements", 2)
add_bullet("1. The system must natively accept standard, widely-used web image formats: .jpg, .jpeg, and .png.")
add_bullet("2. The user interface must proactively restrict input image dimensions to 2048px (for 2x upscaling) or 1024px (for 4x upscaling). This is a hard requirement to prevent the serverless GPUs from experiencing Out-Of-Memory (OOM) crashes.")
add_bullet("3. The API gateway must implement a robust asynchronous queue. This strictly separates the lightweight web server from the heavy GPU tasks, ensuring the site remains responsive even when the AI is processing.")
add_bullet("4. Direct Cloud Upload Optimization: To minimize server memory usage and API latency, the Serverless GPU must upload processed images directly to Cloudinary and return the Cloud URL, rather than passing a massive binary file back through the Node.js API.")
add_bullet("5. The system needs to provide immediate visual feedback, displaying an interactive before-and-after slider of the image.")
add_h("4.4 Cost-Optimized Deployment Architecture", 2)
add_p("A critical engineering constraint of this project was achieving enterprise-grade scalability with minimal initial capital expenditure. The primary Node.js API gateway is hosted using promotional cloud credits from DigitalOcean. The heavy PyTorch inference pipelines rely on Modal.com's base compute allowances, ensuring virtually zero out-of-pocket costs for GPU time during the MVP phase. This financial engineering ensures the SaaS can maintain a high-performance state, handle initial user traction, and validate market fit before transitioning to paid infrastructure tiers.")
doc.add_page_break()

add_h("5. DFD DIAGRAMS", 1)
add_h("Level 0 Data Flow Diagram (Context Diagram)", 2)
add_image_placeholder("Level 0 DFD Image Here")
add_h("Level 1 Data Flow Diagram (Main Architecture)", 2)
add_p("Note: This diagram illustrates the optimized Direct Cloud Upload flow.")
add_image_placeholder("Level 1 DFD Image Here")
add_h("Level 2 Data Flow Diagram (Processing Subsystem)", 2)
add_image_placeholder("Level 2 DFD Image Here")
doc.add_page_break()

add_h("6. ENTITY RELATIONSHIP DIAGRAM (ERD)", 1)
add_p("The database architecture is built on strict relational normalization principles to ensure data integrity. Instead of a separate subscription table, user billing states are tracked directly within the users table, mapped against predefined subscription_plans. The system also maintains comprehensive logging via usage_tracking and analytics tables.")
add_image_placeholder("Database ERD Image Here")
doc.add_page_break()

add_h("7. SNAPSHOTS OF INPUT AND OUTPUT SCREENS", 1)
add_p("To guarantee operational stability under multi-tenant load scenarios, system limitations were stress-tested. The transition from localized CPU processing to GPU environments—combined with direct-to-cloud file uploads—yielded an approximate 96% reduction in payload processing time (averaging 3.2s to 6.8s on serverless APIs compared to >120s on legacy monolithic architectures).")
add_h("Visual Output Results", 2)
add_p("Evaluation of upscaled images assessed human perceptual visual quality specifically concerning texture hallucination.")
add_p("Result 1: Real-ESRGAN Pro (Architectural Artifacts)", bold=True)
add_bullet("Input: Low Resolution, heavily compressed brick wall structure.")
add_bullet("Output: The Real-ESRGAN Pro variant successfully detected structural outlines, eliminating the JPG compression artifacts, and authentically regenerated granular details on the surface without introducing hallucinated color artifacts.")
add_p("Result 2: Real-ESRGAN Anime (Digital Illustration)", bold=True)
add_bullet("Input: Pixelated, jagged contours of vector artwork.")
add_bullet("Output: Running the Real-ESRGAN Anime instance corrected line deviations entirely. The model aggressively isolated the primary contour lines, thickening and sharpening them accurately while smoothing the intermediate color gradients flawlessly, removing pixelated color bleeds.")

add_image_placeholder("React UI Payload Upload Screen (Screenshot)")
add_image_placeholder("Pre vs. Post upscale slider tool showcasing brick texture recovery (Screenshot)")
add_image_placeholder("Processing latency matrix log from API returned payload (Screenshot)")
doc.add_page_break()

add_h("8. CONCLUSION", 1)
add_p("As display technologies advance and screens become increasingly sharper, older digital image data visually degrades, creating a massive commercial demand for intelligent, automated upscaling solutions. This project successfully conceptualized, engineered, and deployed a functioning Software as a Service (SaaS) architecture that rigorously prioritizes backend scalability alongside state-of-the-art visual fidelity.")
add_p("By enforcing a strict separation between the main web routing server and the computationally exhaustive AI inference tasks, the platform maintains high availability. Transitioning away from legacy math-based resizing to advanced neural networks running entirely on serverless GPU clusters solved previous industry challenges regarding slow processing speeds and poor image generation. Furthermore, implementing direct cloud-to-cloud asset transfers vastly reduced Node.js memory bottlenecks. Ultimately, the finalized system consistently processes high-resolution reconstructions in under 10 seconds, proving that developers can reliably integrate standard web APIs with complex, serverless PyTorch ecosystems to build a highly responsive, commercial-ready AI application.")
doc.add_page_break()

add_h("9. LIMITATION", 1)
add_p("While the decoupled serverless architecture works exceptionally well for general use cases, there are a few hard technical limitations inherent to the current deployment ecosystem:")
add_bullet("1. GPU Memory Limits (VRAM Saturation): Image processing via GANs requires loading massive multi-dimensional tensors into the GPU's memory. To stop the serverless NVIDIA T4 GPUs (which are limited to 16GB of VRAM) from crashing due to Out-Of-Memory (OOM) exceptions, the application must strictly cap input dimensions (2048px for 2x, 1024px for 4x).")
add_bullet("2. Cold Start Latency Overheads: Serverless architecture remains cost-effective specifically by spinning down and turning off ephemeral GPU containers when traffic is idle. If a user uploads an image during a \"cold\" phase, the system must provision a new container, initialize the PyTorch environment, and load the heavy model weights into VRAM. This process appends a variable latency overhead (typically 5 to 8 seconds).")
add_bullet("3. Format Support Specificity: Currently, the system is engineered exclusively to process standard, rasterized 2D image matrices (JPEG and PNG). It does not contain the logic to handle vector upsampling, animated .gif sequencing, or temporal .mp4 video frame processing.")
doc.add_page_break()

add_h("10. FUTURE SCOPE", 1)
add_p("In future iterations, the platform's architectural capacities can be expanded:")
add_bullet("1. Video Temporal Upscaling: Upgrading the queueing service and storage architecture to support low-resolution .mp4 payloads, requiring frame-by-frame tensor manipulation and optical flow consistency to prevent flickering.")
add_bullet("2. Dynamic Tensor Chunking Logic: To circumvent the current 16GB VRAM hardware limitations on high-resolution photography, automated image chunking could be implemented. This would mathematically slice a massive image into smaller grids, process those segments concurrently across multiple GPU nodes, and seamlessly stitch them back together.")
add_bullet("3. Dedicated Face Recovery Integration: Utilizing parameterized models like CodeFormer or GFPGAN to target and repair semantic facial geometries independently of the background environment.")
doc.add_page_break()

add_h("11. BIBLIOGRAPHY", 1)
add_p("1. C. Dong, C. C. Loy, K. He, and X. Tang, \"Image Super-Resolution Using Deep Convolutional Networks,\" IEEE Transactions on Pattern Analysis and Machine Intelligence, vol. 38, no. 2, pp. 295-307, Feb. 2016.")
add_p("2. W. Shi et al., \"Real-Time Single Image and Video Super-Resolution Using an Efficient Sub-Pixel Convolutional Neural Network,\" CVPR, 2016.")
add_p("3. C. Ledig et al., \"Photo-Realistic Single Image Super-Resolution Using a Generative Adversarial Network,\" CVPR, 2017.")
add_p("4. X. Wang et al., \"ESRGAN: Enhanced Super-Resolution Generative Adversarial Networks,\" ECCV Workshops, 2018.")
add_p("5. X. Wang, L. Xie, C. Dong, and Y. Shan, \"Real-ESRGAN: Training Real-World Blind Super-Resolution with Pure Synthetic Data,\" ICCV Workshops, 2021.")
add_p("6. Node.js Foundation. (2025). Node.js Documentation. Available: https://nodejs.org/en/docs/.")

doc.save("UpscalePro_Dissertation.docx")
print("SUCCESS: 'UpscalePro_Dissertation.docx' has been created in this folder!")