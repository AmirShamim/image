"""
Shared configuration, styles, and helper functions for the dissertation generator.
All formatting follows IEEE standards with Times New Roman font.
"""
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("CRITICAL: python-docx missing. Run: pip install python-docx")
    sys.exit(1)

# ── Global document instance ──
doc = Document()

# ── Page margins (1 inch = 2.54 cm standard IEEE) ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)   # slightly wider left for binding
    section.right_margin = Cm(2.54)

# ── Default paragraph style ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)


# ═══════════════════════════════════════════════════════
#  HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════

def add_heading_styled(text, level):
    """Add a heading with consistent IEEE-style formatting."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
        elif level == 3:
            run.font.size = Pt(12)
            run.font.bold = True
    return h


def add_para(text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Add a justified paragraph with 1.5 line spacing."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    return p


def add_bullet(text):
    """Add a bullet-point item."""
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p


def add_numbered(text):
    """Add a numbered list item."""
    p = doc.add_paragraph(style='List Number')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p


def add_code_block(text):
    """Add a code snippet in Courier New."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    return p


def add_figure_placeholder(caption_text, fig_number):
    """Add a red placeholder for a screenshot with a proper IEEE figure caption below."""
    # Placeholder box
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(f"\n[ INSERT SCREENSHOT HERE ]\n")
    r.bold = True
    r.font.color.rgb = RGBColor(200, 0, 0)
    r.font.size = Pt(11)
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    rc = cap.add_run(f"Fig. {fig_number}: {caption_text}")
    rc.font.name = 'Times New Roman'
    rc.font.size = Pt(10)
    rc.italic = True
    return cap


def add_figure_explanation(text):
    """Add the explanatory paragraph that follows every figure (teacher requirement)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p


def add_table(headers, rows):
    """Create a formatted table with header row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    doc.add_paragraph()  # spacing after table
    return table


def page_break():
    doc.add_page_break()
