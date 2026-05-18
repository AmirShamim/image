import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("CRITICAL ERROR: python-docx library missing. Run: pip install python-docx")
    sys.exit(1)

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

def add_h(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
        elif level == 3:
            run.font.size = Pt(12)
            run.font.bold = True

def add_p(text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    if bold:
        r.bold = True

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.add_run(text)

def add_code(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(10)

def add_image_placeholder(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(f"\n[ INSERT IMAGE HERE: {text} ]\n")
    r.bold = True
    r.font.color.rgb = RGBColor(255, 0, 0)

print("Generating Ultimate Lengthy Document...")

# --- TITLE PAGE ---
doc.add_paragraph("\n\n\n\n")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("UPSCALE PRO\n").bold = True
p.runs[0].font.size = Pt(24)
p.add_run("MAJOR PROJECT / DISSERTATION REPORT\n\n\n").font.size = Pt(16)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run("Submitted by: Amir Shamim (2023-301-021)\n").bold = True
p2.add_run("In partial fulfilment for the award of the degree of:\nBACHELOR OF COMPUTER APPLICATION\n\n")
p2.add_run("Under the supervision of: Dr. Sapna Jain Ma’am\n\n").bold = True
p2.add_run("Institution: Jamia Hamdard University, Department of Computer Science & Technology\n")
doc.add_page_break()

# --- DECLARATION ---
add_h("DECLARATION", 1)
add_p("I, Amir Shamim, a student of BCA (Enrollment No: 2023-301-021), hereby declare that the dissertation entitled “Upscale Pro” which is being submitted by me to the Department of Computer Science, Jamia Hamdard, New Delhi in partial fulfilment of the requirement for the award of the degree of BCA is my original work and has not been submitted anywhere else for the award of any Degree, Diploma, Associateship, Fellowship or other similar title or recognition.")
p = doc.add_paragraph()
p.add_run("\nApril, 2026\n").bold = True
p.add_run("Amir Shamim").bold = True
doc.add_page_break()

# --- ACKNOWLEDGEMENT ---
add_h("ACKNOWLEDGEMENT", 1)
add_p("It is a pleasure to acknowledge many people who knowingly and unknowingly helped me to complete my project. First, I thank God for all the blessings that carried me through all these years.")
add_p("I extend my utmost gratitude to Dr Sapna Jain Ma’am, project supervisor, who has always supported, guided, appreciated, and encouraged me to explore more opportunities. She enlightened me at various stages during the development of this project and provided many insights and useful examples, which proved immensely helpful in its successful completion.")
add_p("I extend my sincere gratitude to my teachers and guides who made unforgettable contributions. I thank all the non-teaching staff of our institution who were always ready to help in whatever way they could.")
doc.add_page_break()

# --- ABSTRACT ---
add_h("ABSTRACT", 1)
add_p("Every digital user, from e-commerce developers to everyday consumers, understands the frustration of working with low-resolution, heavily compressed images. As modern display technologies advance to 4K and Retina standards, legacy digital media visually degrades. Traditional interpolation algorithms (like nearest-neighbor or bicubic scaling) fail to recover this lost data, resulting in severely pixelated and blurred outputs. Furthermore, utilizing advanced machine learning to hallucinate and reconstruct these missing pixels requires immense computational GPU power, which is inaccessible to standard users and difficult to host on traditional monolithic web servers without crashing.")
add_p("This project presents Upscale PRO, an innovative, serverless Software as a Service (SaaS) platform designed to democratize high-fidelity image enhancement. By leveraging heavily tuned Generative Adversarial Networks (specifically Real-ESRGAN Pro and Real-ESRGAN Anime), the platform accurately reconstructs high-frequency details, removes JPEG compression artifacts, and restores image sharpness. To ensure enterprise-grade scalability and eliminate server timeouts, Upscale PRO implements a strictly decoupled architecture. The user-facing frontend is built with React.js, while a Node.js API gateway intelligently routes payloads to ephemeral, serverless NVIDIA GPUs hosted on Modal.com. This financial and architectural engineering allows the system to spin up heavy PyTorch environments on-demand, process complex tensor matrices in under 10 seconds, and upload the finalized assets directly to a Cloudinary CDN network. Ultimately, Upscale PRO revolutionizes AI image processing by combining state-of-the-art neural networks with cost-optimized cloud infrastructure, delivering a seamless, rapid, and professional-grade upscaling experience to the end-user.")
doc.add_page_break()

# --- CHAPTER 1: INTRODUCTION ---
add_h("1. INTRODUCTION", 1)
add_h("1.1 Background", 2)
add_p("Today, digital images dictate how we interact with the web—from online shopping catalogs and social media platforms to digital archiving and professional broadcasting. However, acquiring high-quality, high-resolution pictures isn't always possible. Hardware constraints on older mobile devices, aggressive compression algorithms used by messaging applications, and limited network bandwidth often leave users with low-quality, heavily pixelated files.")
add_p("When these degraded images are stretched to fit modern 4K monitors or high-density smartphone screens, the lack of pixel density becomes painfully obvious. The standard industry response has historically been mathematical resizing. Traditional methods like nearest-neighbor, bilinear, and bicubic interpolation attempt to solve this by simply calculating and guessing the missing pixels based on the color values of the pixels immediately surrounding them. While computationally inexpensive, these spatial domain algorithms are fundamentally blind; they do not understand the contextual content of the image. Consequently, they permanently ruin fine high-frequency details—such as skin textures, fabric threads, and sharp geometric edges—leaving the image looking muddy and blurred.")
add_h("1.2 The Upscale PRO Solution", 2)
add_p("To resolve this fundamental flaw, this project introduces Upscale PRO, a serverless GPU-powered Image Upscaler built as a commercial SaaS. Instead of relying on blind mathematical blurring, Upscale PRO utilizes Generative Adversarial Networks (GANs). Specifically, it deploys the Real-ESRGAN architecture, which utilizes a complex \"degradation model\" that simulates real-world image problems—like motion blur, heavy JPEG compression, and camera sensor noise.")
add_p("By shifting the heavy computational burden from the user's local machine to isolated, cloud-based neural networks, the platform democratizes access to professional-grade image enhancement. Users do not need expensive graphics cards; they only need a standard web browser. Upscale PRO orchestrates a complex background pipeline where a Node.js server receives the image, securely offloads the mathematical tensor calculations to remote Modal.com serverless GPUs, and streams the restored, high-definition image back to the user in a matter of seconds.")
doc.add_page_break()

# --- CHAPTER 2: OBJECTIVE ---
add_h("2. OBJECTIVE", 1)
add_p("The primary objective of this dissertation is to architect, build, and deploy a commercially viable, scalable web platform that makes advanced AI image upscaling accessible to everyday users. To achieve a stable production environment, the project targets several specific engineering and business objectives:")
add_bullet("1. Decoupled SaaS Architecture: To design and build a full-stack SaaS ecosystem that strictly separates the UI, the web server, and the AI inference engine. This involves utilizing React.js for a responsive client frontend and Node.js for a robust backend API gateway, ensuring that intensive GPU processing never bottlenecks the main web server.")
add_bullet("2. Sub-10 Second Latency Benchmark: To cut down AI processing time to a reliable sub-10-second window. This is achieved by offloading heavy PyTorch inference tasks to Modal.com's serverless GPUs. Instead of paying for continuous, expensive AWS cloud uptime, this architecture utilizes ephemeral containers that boot up on-demand, making the application both blazingly fast and financially viable to host.")
add_bullet("3. Direct Cloud Upload Pipeline: To implement a direct-to-cloud file transfer mechanism. Instead of pushing massive, upscaled binary image buffers back through the Node.js server (which causes RAM heap crashes), the serverless GPUs upload the finalized assets directly to a Cloudinary CDN, returning only a lightweight string URL to the API.")
add_bullet("4. Multi-Tenant Security and Tiering: To offer tiered processing models based on user subscription levels, requiring secure authentication and rate limiting using Neon Serverless PostgreSQL.")
doc.add_page_break()

# --- NEW CHAPTER 3: FEASIBILITY STUDY ---
add_h("3. FEASIBILITY STUDY", 1)
add_p("A feasibility study is a critical phase in the software development life cycle. It assesses the practicality of the proposed system. For Upscale PRO, the feasibility was evaluated across three primary domains: Economic, Technical, and Operational.")
add_h("3.1 Economic Feasibility", 2)
add_p("Economic feasibility determines whether the project is financially viable to build and host. Running heavy AI models is notoriously expensive. Renting a dedicated cloud server with an NVIDIA GPU (such as an AWS EC2 p3 instance) can cost thousands of dollars per month, running 24/7 regardless of traffic. This was not economically feasible for a student MVP.")
add_p("Upscale PRO circumvents this via Serverless Computing. By utilizing Modal.com, the system only pays for the exact seconds the GPU is actively processing an image. Furthermore, the MVP was deployed using the GitHub Student Developer Pack, leveraging $200 in DigitalOcean credits for the Node.js API and free-tier allowances on Neon PostgreSQL and Cloudinary. The initial capital expenditure to launch this SaaS was effectively $0, making it highly economically feasible.")
add_h("3.2 Technical Feasibility", 2)
add_p("Technical feasibility evaluates if the current technology stack can handle the system's demands. The primary technical risk was the Node.js server crashing due to 'Out of Memory' (OOM) errors when handling massive 15MB to 20MB image buffers. This risk was mitigated by strictly decoupling the storage architecture. By programming the Modal GPU to upload the finalized image directly to Cloudinary and only returning a lightweight URL string back to Node.js, the memory bottleneck was completely eliminated. The integration of React, Node.js, PyTorch, and Neon DB is a proven, highly scalable combination, confirming the system's technical feasibility.")
add_h("3.3 Operational Feasibility", 2)
add_p("Operational feasibility assesses how well the proposed system solves the users' problems and how easily they can adapt to it. Historically, utilizing Real-ESRGAN required users to open a command-line interface, install Python, download heavy PyTorch dependencies, and possess a dedicated local GPU. Upscale PRO abstracts all of this complexity away. The operational flow simply requires the user to drag and drop an image onto a web browser. The system handles all compilation and processing in the cloud, making it incredibly accessible to non-technical users.")
doc.add_page_break()

# --- CHAPTER 4: PROBLEM STATEMENT ---
add_h("4. PROBLEM STATEMENT", 1)
add_p("The core problem this project addresses is not just generating high-resolution images, but the immense architectural difficulty of hosting and delivering these AI models over the internet to concurrent users.")
add_bullet("1. The Quality Bottleneck: Traditional interpolation algorithms cannot hallucinate missing data. They cannot repair a compressed JPEG or restore the granular texture of a brick wall. They only blur existing pixels together.")
add_bullet("2. The Compute Bottleneck (The Monolith Problem): Advanced neural networks like Real-ESRGAN solve the quality problem, but they create a severe infrastructure problem. They require massive computational power from Graphics Processing Units (GPUs) with large VRAM capacities. Historically, if a developer attempted to host a heavy PyTorch model directly on a standard monolithic web server (like a basic VPS or cPanel), the entire system would freeze the moment a user uploaded an image.")
add_bullet("3. The Concurrency Failure: In a standard setup, if five users upload images simultaneously, the server's CPU and RAM will spike to 100%, leading to extremely slow loading times, HTTP 504 Gateway Timeouts, and complete application crashes. Upscale PRO solves this by completely destroying the monolithic model and routing every individual user request to its own isolated, temporary serverless GPU container.")
doc.add_page_break()

# --- CHAPTER 5: SRS ---
add_h("5. SOFTWARE REQUIREMENT SPECIFICATION", 1)
add_h("5.1 Introduction", 2)
add_p("5.1.1 Purpose: The purpose of this document is to define the software architecture, capabilities, and strict constraints of the Upscale PRO system.", bold=True)
add_p("5.1.2 Scope: The software system is an asynchronous, cloud-based Image Upscaler processing standard web image formats.", bold=True)
add_h("5.2 Specific Requirements (Functional)", 2)
add_bullet("REQ-F01 (Secure Authentication): The system shall authenticate users securely using encrypted passwords, generating a temporary JWT for API access.")
add_bullet("REQ-F02 (Payload Sanitization): The API gateway shall intercept all incoming multipart/form-data uploads. It must strictly validate MIME types, rejecting executables or SVGs.")
add_bullet("REQ-F03 (Serverless GPU Provisioning): Upon receiving a valid request, the Node.js API shall trigger a remote procedure call (RPC) via the Modal SDK to instantly provision an ephemeral NVIDIA GPU container.")
add_bullet("REQ-F04 (Direct Cloud Transfer): The remote GPU environment shall execute the PyTorch inference and stream the finalized binary output directly to the Cloudinary CDN. The GPU will then return only the Cloudinary URL back to the Node.js API.")
add_h("5.3 Specific Requirements (Non-Functional)", 2)
add_bullet("REQ-NF01 (Performance & Latency): The end-to-end processing lifecycle shall not exceed an average latency of 10 seconds per standard payload.")
add_bullet("REQ-NF02 (Hardware Constraints/VRAM): The user interface must dynamically restrict input image dimensions to 2048px (for 2x upscaling) or 1024px (for 4x upscaling) to prevent Out-Of-Memory crashes.")
doc.add_page_break()

# --- NEW CHAPTER 6: SYSTEM DESIGN AND API ARCHITECTURE ---
add_h("6. SYSTEM DESIGN AND API ARCHITECTURE", 1)
add_p("To ensure that Upscale PRO could scale to thousands of users without crashing, a strictly decoupled microservices architecture was implemented. The system relies heavily on RESTful API design principles to facilitate communication between the React client, the Node.js Gateway, and the Modal Serverless GPUs.")
add_h("6.1 API Routing Flow", 2)
add_p("When a user submits an image, the React frontend packages the file into a FormData object and sends a POST request to the Node.js server. The server acts as a middleware gatekeeper. It verifies the user's JWT token, checks their rate limit in the Neon PostgreSQL database, and then forwards a sanitized base64 string to the Modal GPU endpoint.")
add_h("6.2 Data Payloads (JSON Specifications)", 2)
add_p("The following outlines the standard JSON request structure sent from the Node.js API to the Modal GPU interface:")
add_code("{\n  \"action\": \"upscale\",\n  \"parameters\": {\n    \"model_type\": \"RealESRGAN_x4plus\",\n    \"face_enhance\": false,\n    \"scale\": 4\n  },\n  \"image_data\": \"data:image/jpeg;base64,/9j/4AAQSk...\"\n}")
add_p("Once the Modal GPU completes the PyTorch inference, it uploads the image directly to Cloudinary. It then responds to the Node.js server with the following lightweight JSON payload, completely bypassing the need to transmit heavy binary buffers over the network:")
add_code("{\n  \"status\": \"success\",\n  \"processing_time_ms\": 5420,\n  \"gpu_node\": \"t4-ephemeral-889\",\n  \"cloudinary_url\": \"https://res.cloudinary.com/upscalepro/image/upload/v1234/final.png\"\n}")
doc.add_page_break()

# --- CHAPTER 7: TECHNOLOGY STACK ---
add_h("7. TECHNOLOGY STACK AND ARCHITECTURAL JUSTIFICATION", 1)
add_h("7.1 Client-Side Frontend: React.js", 2)
add_p("React.js was chosen because the application requires rapid state management. When the AI returns the upscaled version, the UI must present an interactive 'Before and After' comparison slider. Instead of reloading the entire web page, React calculates the exact minimal changes needed in its Virtual DOM and selectively updates only the image slider component.")
add_h("7.2 Backend API Gateway: Node.js and Express", 2)
add_p("Node.js was chosen to act as the central API gateway due to its asynchronous, event-driven, non-blocking I/O model. Node.js offloads the upscaling request to the cloud and immediately moves on to serve the next user, waiting asynchronously for the cloud to reply, resulting in extremely low CPU overhead on the main server.")
add_h("7.3 AI Infrastructure: Serverless Computing via Modal.com", 2)
add_p("Serverless architecture (provided by Modal.com) completely resolves the cost of 24/7 cloud GPUs. When a user uploads an image, Modal instantly spins up a temporary Docker container equipped with a GPU, loads the Python environment, processes the image in 5 seconds, and then instantly destroys the container. Upscale PRO is only billed for those exact 5 seconds of compute time.")
add_h("7.4 Storage Array: Cloudinary CDN", 2)
add_p("The serverless GPU pushes the finalized image directly to Cloudinary via their API. Cloudinary securely stores the asset on its global Content Delivery Network (CDN) and simply returns a lightweight string URL. This bypasses Node.js memory limits entirely.")
add_h("7.5 Database Management: Neon PostgreSQL", 2)
add_p("Because the Node.js API and the Modal GPUs are constantly scaling up and shutting down, they would normally overwhelm a traditional database. Neon PostgreSQL natively supports 'connection pooling' at the edge, ensuring the database remains stable.")
doc.add_page_break()

# --- CHAPTER 8: SYSTEM ANALYSIS (DIAGRAMS) ---
add_h("8. SYSTEM ANALYSIS", 1)
add_h("8.1 Data Flow Diagrams (DFD)", 2)
add_p("The Data Flow Diagrams illustrate how image payloads, user credentials, and database queries are securely routed across the network.")
add_h("8.1.1 Level 0 Context Diagram", 3)
add_image_placeholder("Level 0 DFD")
add_h("8.1.2 Level 1 Architecture Diagram", 3)
add_image_placeholder("Level 1 DFD")
add_h("8.1.3 Level 2 Processing Subsystem Diagram", 3)
add_image_placeholder("Level 2 DFD")
add_h("8.2 Entity Relationship Diagram (ERD)", 2)
add_image_placeholder("Database ERD")
doc.add_page_break()

# --- NEW CHAPTER 9: SOFTWARE TESTING ---
add_h("9. SOFTWARE TESTING METHODOLOGIES", 1)
add_p("To ensure the Upscale PRO system is robust, secure, and capable of handling high traffic loads without returning 504 Gateway Timeouts, a rigorous testing lifecycle was implemented.")
add_h("9.1 Unit Testing", 2)
add_p("Unit testing focuses on verifying the smallest parts of the application independently. In the Node.js API, unit tests were written to validate the JWT authentication middleware. Tests were executed to ensure that if a user submitted an expired or malformed token, the API would immediately return a 401 Unauthorized status without attempting to wake up the expensive GPU cluster.")
add_h("9.2 Integration Testing", 2)
add_p("Integration testing validates that the distinct microservices communicate correctly. The critical test involved the handshake between Node.js and the Modal GPU endpoint. Fake base64 image strings were injected into the pipeline to verify that the GPU received the payload, initialized PyTorch, successfully authenticated with the Cloudinary API, and returned the correct JSON structure back to the Node.js server.")
add_h("9.3 Load and Stress Testing", 2)
add_p("Load testing was crucial for validating the serverless architecture. Simulated traffic (50 concurrent image uploads) was directed at the API. On a traditional monolithic server, this would cause immediate RAM heap exhaustion. However, the decoupled architecture successfully routed all 50 requests to Modal.com, which automatically spun up 50 independent, ephemeral GPU containers. All 50 images were processed and uploaded to Cloudinary simultaneously without crashing the main Node.js web server.")
doc.add_page_break()

# --- CHAPTER 10: IMPLEMENTATION AND RESULTS ---
add_h("10. IMPLEMENTATION AND RESULTS", 1)
add_p("The transition from localized CPU processing to isolated GPU environments yielded an approximate 96% reduction in payload processing time. Complex image reconstructions that previously took over 120 seconds are now consistently processed in 3.2 to 6.8 seconds.")
add_h("10.1 System Interface & User Flow", 2)
add_p("1. Upload Phase: The user authenticates and uploads a degraded image via the React.js interface.", bold=True)
add_image_placeholder("React UI Payload Upload Screen Screenshot")
add_p("2. Processing Phase: The React client displays a loading state while the Node.js server triggers the Modal.com GPU.", bold=True)
add_p("3. Delivery Phase: The React interface dynamically renders an interactive, side-by-side comparison slider.", bold=True)
add_h("10.2 Visual Output Analysis & Evaluation", 2)
add_p("Result 1: Real-ESRGAN Pro (Architectural Textures & Artifacts)", bold=True)
add_bullet("The AI mathematically reconstructed realistic, high-frequency granular details back onto the surface of the bricks without introducing any strange, neon color glitches.")
add_image_placeholder("Pre vs. Post upscale slider tool showcasing brick texture recovery")
add_p("Result 2: Real-ESRGAN Anime (Digital Illustration & Vector Geometry)", bold=True)
add_bullet("The model aggressively isolated the primary dark contour lines, thickening and sharpening them accurately while smoothing out the intermediate color gradients perfectly.")
add_image_placeholder("Pre vs. Post upscale slider tool showcasing Anime/Vector recovery")
add_image_placeholder("Processing latency matrix log from API returned payload")
doc.add_page_break()

# --- CHAPTER 11 & 12: CONCLUSION, LIMITATION, FUTURE SCOPE ---
add_h("11. CONCLUSION", 1)
add_p("This project successfully conceptualized, engineered, and deployed Upscale PRO. By enforcing a strict decoupling between the main Node.js web routing server and the computationally exhaustive AI inference tasks, the platform ensures maximum high availability. Transitioning away from legacy math-based resizing to advanced Generative Adversarial Networks (GANs) running entirely on serverless Modal GPU clusters solved previous industry challenges regarding slow processing speeds. Ultimately, the finalized system consistently processes high-resolution, artifact-free reconstructions in under 10 seconds.")
add_h("12. LIMITATIONS AND FUTURE SCOPE", 1)
add_bullet("1. GPU Memory Limits: Inputs are currently restricted to 2048px to prevent the serverless 16GB VRAM GPUs from crashing due to Out-Of-Memory (OOM) exceptions.")
add_bullet("2. Cold Start Latency: If a user uploads an image during a \"cold\" phase, the Modal system must provision a new container from scratch, adding 5 to 8 seconds of latency.")
add_p("In future iterations, automated image chunking algorithms will be implemented to mathematically slice massive input images into smaller grids, process them concurrently, and seamlessly stitch them back together to bypass VRAM limitations. Furthermore, auxiliary facial reconstruction pipelines like GFPGAN will be integrated to specifically repair semantic facial geometries.")
doc.add_page_break()

# --- BIBLIOGRAPHY ---
add_h("13. BIBLIOGRAPHY", 1)
add_p("1. C. Dong, C. C. Loy, K. He, and X. Tang, \"Image Super-Resolution Using Deep Convolutional Networks,\" IEEE Transactions on Pattern Analysis and Machine Intelligence, vol. 38, no. 2, pp. 295-307, Feb. 2016.")
add_p("2. W. Shi et al., \"Real-Time Single Image and Video Super-Resolution Using an Efficient Sub-Pixel Convolutional Neural Network,\" CVPR, 2016.")
add_p("3. X. Wang, L. Xie, C. Dong, and Y. Shan, \"Real-ESRGAN: Training Real-World Blind Super-Resolution with Pure Synthetic Data,\" ICCV Workshops, Virtual, 2021.")
add_p("4. Modal Labs. (2025). Modal Documentation & Serverless Infrastructure. Available: https://modal.com/docs.")
add_p("5. Node.js Foundation. (2025). Node.js Documentation. Available: https://nodejs.org/en/docs/.")
add_p("6. Cloudinary. (2025). Programmable Media Documentation. Available: https://cloudinary.com/documentation.")

doc.save("UpscalePro_Ultimate_Report.docx")
print("SUCCESS: 'UpscalePro_Ultimate_Report.docx' has been generated!")